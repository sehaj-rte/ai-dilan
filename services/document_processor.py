import os
import logging
from typing import Dict, Any, List
import PyPDF2
import docx
import csv
import io
from PIL import Image
import pytesseract
import speech_recognition as sr
from pydub import AudioSegment

logger = logging.getLogger(__name__)

class DocumentProcessor:
    """Service to extract text content from various file types"""
    
    def __init__(self):
        self.supported_types = {
            'application/pdf': self._extract_pdf_text,
            'application/msword': self._extract_doc_text,
            'application/vnd.openxmlformats-officedocument.wordprocessingml.document': self._extract_docx_text,
            'text/plain': self._extract_text_file,
            'text/csv': self._extract_csv_text,
            'image/jpeg': self._extract_image_text,
            'image/png': self._extract_image_text,
            'image/gif': self._extract_image_text,
            'audio/mpeg': self._extract_audio_text,
            'audio/wav': self._extract_audio_text,
            'video/mp4': self._extract_video_text,
            'video/avi': self._extract_video_text
        }
    
    def extract_text(self, file_content: bytes, content_type: str, filename: str) -> Dict[str, Any]:
        """Extract text from file content based on content type"""
        try:
            print(f"\U0001f4d1 Document Processor: Extracting text from {filename} ({content_type})")
            if content_type not in self.supported_types:
                print(f"\U0001f6ab Document Processor: Unsupported file type {content_type}")
                return {
                    "success": False,
                    "error": f"Unsupported file type: {content_type}"
                }
            
            extractor = self.supported_types[content_type]
            text_content = extractor(file_content, filename)
            
            if not text_content or len(text_content.strip()) == 0:
                print(f"\U0001f6ab Document Processor: No text content extracted from {filename}")
                return {
                    "success": False,
                    "error": "No text content could be extracted from the file"
                }
            
            print(f"\U0001f389 Document Processor: Successfully extracted {len(text_content)} characters from {filename}")
            # Extract enhanced metadata
            metadata = self._extract_metadata(text_content, content_type, filename, file_content)
            
            return {
                "success": True,
                "text": text_content,
                "content_type": content_type,
                "filename": filename,
                "word_count": len(text_content.split()),
                "metadata": metadata
            }
            
        except Exception as e:
            logger.error(f"Error extracting text from {filename}: {str(e)}")
            return {
                "success": False,
                "error": f"Text extraction failed: {str(e)}"
            }
    
    def _extract_pdf_text(self, file_content: bytes, filename: str) -> str:
        """Extract text from PDF files"""
        try:
            pdf_file = io.BytesIO(file_content)
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            
            text = ""
            for page in pdf_reader.pages:
                text += page.extract_text() + "\n"
            
            return text.strip()
        except Exception as e:
            logger.error(f"PDF extraction error for {filename}: {str(e)}")
            return ""
    
    def _extract_docx_text(self, file_content: bytes, filename: str) -> str:
        """Extract text from DOCX files"""
        try:
            doc_file = io.BytesIO(file_content)
            doc = docx.Document(doc_file)
            
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"
            
            return text.strip()
        except Exception as e:
            logger.error(f"DOCX extraction error for {filename}: {str(e)}")
            return ""
    
    def _extract_doc_text(self, file_content: bytes, filename: str) -> str:
        """Extract text from DOC files (basic implementation)"""
        try:
            # For .doc files, we'll return a placeholder for now
            # In production, you might want to use python-docx2txt or antiword
            logger.warning(f"DOC file processing not fully implemented for {filename}")
            return f"DOC file content from {filename} - requires additional processing"
        except Exception as e:
            logger.error(f"DOC extraction error for {filename}: {str(e)}")
            return ""
    
    def _extract_text_file(self, file_content: bytes, filename: str) -> str:
        """Extract text from plain text files"""
        try:
            # Try UTF-8 first, then fallback to other encodings
            encodings = ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1']
            
            for encoding in encodings:
                try:
                    return file_content.decode(encoding)
                except UnicodeDecodeError:
                    continue
            
            # If all encodings fail, return error
            raise UnicodeDecodeError("Unable to decode text file with any supported encoding")
            
        except Exception as e:
            logger.error(f"Text file encoding error for {filename}: {str(e)}")
            return ""
    
    def _extract_csv_text(self, file_content: bytes, filename: str) -> str:
        """Extract text from CSV files"""
        try:
            csv_text = file_content.decode('utf-8')
            csv_file = io.StringIO(csv_text)
            reader = csv.reader(csv_file)
            
            text = ""
            for row_num, row in enumerate(reader):
                if row_num == 0:
                    # Add headers with context
                    text += "Headers: " + " | ".join(row) + "\n"
                else:
                    text += " | ".join(row) + "\n"
            
            return text.strip()
        except Exception as e:
            logger.error(f"CSV extraction error for {filename}: {str(e)}")
            return ""
    
    def _extract_image_text(self, file_content: bytes, filename: str) -> str:
        """Extract text from images using OCR"""
        try:
            # Check if tesseract is available
            if not self._is_tesseract_available():
                logger.warning("Tesseract OCR not available - skipping image text extraction")
                return f"Image file {filename} - OCR not available"
            
            image = Image.open(io.BytesIO(file_content))
            text = pytesseract.image_to_string(image)
            
            return text.strip()
        except Exception as e:
            logger.error(f"Image OCR error for {filename}: {str(e)}")
            return f"Image file {filename} - OCR failed: {str(e)}"
    
    def _extract_audio_text(self, file_content: bytes, filename: str) -> str:
        """Extract text from audio files using speech recognition"""
        try:
            # Create temporary file for audio processing
            temp_audio_path = f"/tmp/{filename}"
            
            with open(temp_audio_path, 'wb') as f:
                f.write(file_content)
            
            # Convert to WAV if needed
            if filename.lower().endswith('.mp3'):
                audio = AudioSegment.from_mp3(temp_audio_path)
                wav_path = temp_audio_path.replace('.mp3', '.wav')
                audio.export(wav_path, format="wav")
                temp_audio_path = wav_path
            
            # Use speech recognition
            recognizer = sr.Recognizer()
            with sr.AudioFile(temp_audio_path) as source:
                audio_data = recognizer.record(source)
                text = recognizer.recognize_google(audio_data)
            
            # Clean up temporary files
            if os.path.exists(temp_audio_path):
                os.remove(temp_audio_path)
            
            return text.strip()
            
        except Exception as e:
            logger.error(f"Audio transcription error for {filename}: {str(e)}")
            return f"Audio file {filename} - transcription failed: {str(e)}"
    
    def _extract_video_text(self, file_content: bytes, filename: str) -> str:
        """Extract text from video files (placeholder)"""
        try:
            # Video processing would require additional libraries like moviepy
            # For now, return a placeholder
            logger.warning(f"Video processing not implemented for {filename}")
            return f"Video file {filename} - processing not yet implemented"
        except Exception as e:
            logger.error(f"Video processing error for {filename}: {str(e)}")
            return ""
    
    def _is_tesseract_available(self) -> bool:
        """Check if Tesseract OCR is available on the system"""
        try:
            pytesseract.get_tesseract_version()
            return True
        except Exception:
            return False
    
    def _extract_metadata(self, text_content: str, content_type: str, filename: str, file_content: bytes) -> Dict[str, Any]:
        """Extract enhanced metadata from document"""
        try:
            metadata = {
                "document_type": self._detect_document_type(content_type),
                "language": self._detect_language(text_content),
                "page_count": self._get_page_count(content_type, file_content),
                "has_images": self._has_images(content_type, file_content),
                "has_tables": self._has_tables(text_content),
                "extracted_text_preview": text_content[:500] + "..." if len(text_content) > 500 else text_content
            }
            return metadata
        except Exception as e:
            logger.error(f"Error extracting metadata from {filename}: {str(e)}")
            return {}
    
    def _detect_document_type(self, content_type: str) -> str:
        """Detect document type from content type"""
        type_mapping = {
            'application/pdf': 'pdf',
            'application/msword': 'doc',
            'application/vnd.openxmlformats-officedocument.wordprocessingml.document': 'docx',
            'text/plain': 'txt',
            'text/csv': 'csv',
            'image/jpeg': 'image',
            'image/png': 'image',
            'image/gif': 'image',
            'audio/mpeg': 'audio',
            'audio/wav': 'audio',
            'video/mp4': 'video',
            'video/avi': 'video'
        }
        return type_mapping.get(content_type, 'unknown')
    
    def _detect_language(self, text: str) -> str:
        """Detect language of the text (basic implementation)"""
        try:
            # Simple language detection based on common words
            # In production, you might want to use langdetect or similar library
            english_words = ['the', 'and', 'is', 'in', 'to', 'of', 'a', 'that', 'it', 'with']
            text_lower = text.lower()
            english_count = sum(1 for word in english_words if word in text_lower)
            
            if english_count >= 3:
                return 'en'
            else:
                return 'unknown'
        except Exception:
            return 'unknown'
    
    def _get_page_count(self, content_type: str, file_content: bytes) -> int:
        """Get page count for PDF documents"""
        try:
            if content_type == 'application/pdf':
                pdf_file = io.BytesIO(file_content)
                pdf_reader = PyPDF2.PdfReader(pdf_file)
                return len(pdf_reader.pages)
            return 1
        except Exception:
            return 1
    
    def _has_images(self, content_type: str, file_content: bytes) -> bool:
        """Check if document contains images"""
        try:
            if content_type.startswith('image/'):
                return True
            # For other document types, this would require more complex analysis
            return False
        except Exception:
            return False
    
    def _has_tables(self, text: str) -> bool:
        """Check if document contains tables (basic detection)"""
        try:
            # Simple table detection based on common patterns
            table_indicators = ['|', '\t\t', '   ', 'Table', 'Column', 'Row']
            return any(indicator in text for indicator in table_indicators)
        except Exception:
            return False

# Create singleton instance
document_processor = DocumentProcessor()
